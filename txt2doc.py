from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from ultralytics.utils.plotting import colors
from tqdm import tqdm
import chardet
import os.path as osp
import random
f = open('ppocr_keys_v1.txt', 'r', encoding='utf-8')
char_dict = [line.strip() for line in f.readlines() if line.strip()]
f.close()

def is_valid_str(line):
    if len(line)>0:
        # 是不是都在char_dict里
        ischar = 0
        for char in line:
            if char in char_dict:
                ischar += 1
        if ischar/len(line)>0.5: # 50%以上不是乱码
            return True
    return False

字体列表 = [
    '微软雅黑', 
    '华文细黑', 
    '华文黑体', 
    '华文楷体', 
    '华文细宋', 
    '华文宋体', 
]
行间距 = [1, 1.5, 2]
def detect_encoding(file_path):
    # 以二进制模式打开文件，读取文件内容
    with open(file_path, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)  # 使用chardet检测编码
        return result['encoding']
def txt2doc(txt_path, doc_path, doc_length=1000):
    encoding = detect_encoding(txt_path)
    try:
        f = open(txt_path, 'r', encoding=encoding, errors='ignore')
        lines = f.read().split('\n')
        f.close()
    except:
        return txt_path
    txt_path = Path(txt_path)
    doc_path = Path(doc_path)
    if doc_path.exists():
        # 之前已经处理过了
        return
    # if len(lines)> doc_length:
    #     # 分文件夹保存
    #     doc_dir = doc_path.parent.joinpath(doc_path.stem)
    # else:
        # 直接使用文件
    doc_dir = doc_path.parent
    doc_dir.mkdir(parents=True, exist_ok=True)
    # last_line_no = 0
    # log_path = doc_dir.joinpath('log.txt')
    # if log_path.exists():
    #     with open(log_path, 'r', encoding='utf-8') as f:
    #         log_lines = f.readlines()
    #         if len(log_lines):
    #             last_line_no = int(log_lines[-1])

    # 创建一个新的Word文档
    # start = (last_line_no // doc_length) * doc_length
    start = 0
    end = start + doc_length if start+doc_length<len(lines) else len(lines)
    doc_path = doc_dir.joinpath(f"{txt_path.stem}_{start}_{end}.docx")
    if doc_path.exists():
        return
    line_num = 0
    bar = tqdm(total=end, desc=f'{txt_path.stem}')

    doc = Document()
    invalid_num = 0
    for j, line in enumerate(lines):
        bar.update(1)
        if invalid_num>20:
            return txt_path
        if not is_valid_str(line):
            invalid_num += 1
            continue
        if j < start:
            continue
        # 添加一个段落并获取其运行对象
        paragraph = doc.add_paragraph()
        # log = open(log_path, 'w', encoding='utf-8')
        # log.write(f'{j}\n')
        # log.close()
        try:
            run = paragraph.add_run(line)
        except:
            continue
        # 设置字体名称
        run.font.name = random.choice(字体列表)
        # 设置字体大小
        run.font.size = Pt(random.randint(8,20))  # 设置字体为14磅
        # 设置字体颜色为红色
        b, g, r = colors(random.randint(0, 20), False)
        if min(b, g, r) > 200:
            b, g, r = min(b, 230), min(g, 230), min(r, 230)
        # 设置字体颜色,一半概率以上是黑色
        if random.random() > 0.5:
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run.font.color.rgb = RGBColor(r, g, b)  # 红色
        # 如果需要设置中文字体（例如微软雅黑），需要额外处理
        r = run._element
        rPr = r.get_or_add_rPr()
        eastAsia = OxmlElement('w:eastAsia')
        eastAsia.set(qn('w:val'), random.choice(字体列表))
        rPr.append(eastAsia)
        # 设置段落的行间距（可以根据需要调整）
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = random.choice(行间距)  # 设置行间距为1.5倍行距
        # 随机设置为居中对齐, 0表示左对齐, 1表示居中对齐, 2表示右对齐
        paragraph_format.alignment = random.randint(0,2) # 1表示居中对齐
        line_num += 1
        bar.set_description(f'{txt_path.stem} : {line_num:04d}, {line}')
        if line_num==end or j>=len(lines)-1:
            # 保存文档
            bar.close()
            bar = tqdm(total=end, desc=f'{Path(txt_path).stem}')
            doc_path = doc_dir.joinpath(f"{txt_path.stem}_{start}_{end}.docx")
            doc.save(doc_path)
            if j>=len(lines)-1:
                break
            start = end
            end = start + doc_length if len(lines) >= doc_length else len(lines)
    # doc_path = doc_dir.joinpath(f"{txt_path.stem}_{start}_{line_num}.docx")
    # doc.save(doc_path)
    print(f'save in {doc_path}')
if __name__ == '__main__':
    txts = Path(r'D:\SCJT\data\pdf\小说\txts').rglob('*.txt')
    error_txts = []
    for txt in txts:
        # if not '边荒传说' in str(txt):
        #     continue
        tard = str(txt).split('txts')[0]
        tard = osp.join(tard, 'docxs')
        tar = osp.join(tard, txt.stem + '.docx')
        if not osp.exists(tar):
            t = txt2doc(str(txt), tar)
            if t:
                error_txts.append(t)
    for t in error_txts:
        print(t)
            # print(tar)
    # src = r'D:\SCJT\data\pdf\小说\txts\uichuideng_downcc.com\鬼吹灯1全本.txt'
    # tar = src.replace('txts', 'docxs').replace('.txt', '.docx')
    # txt2doc(src, tar)
# D:\SCJT\data\pdf\小说\txts